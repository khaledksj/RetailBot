import os
import subprocess
import tempfile
import regex as re
import fitz  # PyMuPDF
import mammoth
from docx import Document
from pathlib import Path
import logging

logger = logging.getLogger(__name__)

ARABIC_RE = re.compile(r'\p{Arabic}', re.UNICODE)

def contains_arabic(s: str, min_ratio=0.15) -> bool:
    """Check if text contains sufficient Arabic content."""
    if not s: 
        return False
    arab = len(ARABIC_RE.findall(s))
    return (arab / max(1, len(s))) >= min_ratio

def normalize_arabic(s: str) -> str:
    """Normalize Arabic text by removing diacritics and unifying variants."""
    if not s:
        return ""
    
    # Remove diacritics and tatweel
    s = re.sub(r'[\u064B-\u065F]', '', s)  # Remove diacritics
    s = s.replace('\u0640', '')  # Remove tatweel
    
    # Unify Arabic letter variants
    s = (s.replace('أ','ا').replace('إ','ا').replace('آ','ا')
           .replace('ى','ي').replace('ٱ','ا').replace('ؤ','و').replace('ئ','ي'))
    
    # Normalize Arabic-Indic digits to ASCII
    s = s.translate(str.maketrans('٠١٢٣٤٥٦٧٨٩', '0123456789'))
    s = s.translate(str.maketrans('۰۱۲۳۴۵۶۷۸۹', '0123456789'))
    
    # Clean up whitespace
    return re.sub(r'\s+', ' ', s).strip()

def _pdf_text(path: str) -> str:
    """Extract text from PDF using PyMuPDF."""
    text = []
    try:
        with fitz.open(path) as doc:
            for page in doc:
                page_text = page.get_text("text")
                if page_text:
                    text.append(page_text)
        return "\n".join(text)
    except Exception as e:
        logger.error(f"PyMuPDF extraction failed: {e}")
        return ""

def _ocr_pdf(path_in: str, langs="ara+eng") -> str:
    """Perform OCR on PDF with Arabic language support."""
    try:
        with tempfile.TemporaryDirectory() as td:
            out_pdf = os.path.join(td, "ocr.pdf")
            
            # Check if tesseract is available
            result = subprocess.run(["tesseract", "--version"], 
                                 capture_output=True, text=True)
            if result.returncode != 0:
                logger.warning("Tesseract not available, skipping OCR")
                return ""
            
            # Try OCR with Arabic + English
            try:
                subprocess.run([
                    "tesseract", path_in, out_pdf.replace('.pdf', ''),
                    "-l", langs, "--psm", "6"
                ], check=True, capture_output=True)
                
                # Read the resulting text file
                txt_file = out_pdf.replace('.pdf', '.txt')
                if os.path.exists(txt_file):
                    with open(txt_file, 'r', encoding='utf-8') as f:
                        return f.read()
            except subprocess.CalledProcessError as e:
                logger.warning(f"Tesseract OCR failed: {e}")
                
        return ""
    except Exception as e:
        logger.error(f"OCR process failed: {e}")
        return ""

def _convert_to_pdf(path_in: str) -> str:
    """Convert document to PDF using LibreOffice."""
    try:
        out_dir = tempfile.mkdtemp()
        
        # Check if LibreOffice is available
        result = subprocess.run(["soffice", "--version"], 
                               capture_output=True, text=True)
        if result.returncode != 0:
            logger.warning("LibreOffice not available, skipping conversion")
            return ""
        
        # Convert to PDF
        subprocess.run([
            "soffice", "--headless", "--convert-to", "pdf", 
            "--outdir", out_dir, path_in
        ], check=True, capture_output=True, timeout=60)
        
        # Find the generated PDF
        base = os.path.splitext(os.path.basename(path_in))[0] + ".pdf"
        pdf_path = os.path.join(out_dir, base)
        
        if os.path.exists(pdf_path):
            return pdf_path
        else:
            logger.error(f"PDF conversion failed - output file not found: {pdf_path}")
            return ""
            
    except subprocess.TimeoutExpired:
        logger.error("LibreOffice conversion timed out")
        return ""
    except subprocess.CalledProcessError as e:
        logger.error(f"LibreOffice conversion failed: {e}")
        return ""
    except Exception as e:
        logger.error(f"Document conversion failed: {e}")
        return ""

def _read_docx(path: str) -> str:
    """Extract text from DOCX using multiple methods."""
    try:
        # Try mammoth first for better formatting
        with open(path, "rb") as f:
            result = mammoth.convert_to_html(f)
            html = result.value
            
        # Strip HTML tags
        txt = re.sub(r'<[^>]+>', ' ', html)
        txt = re.sub(r'\s+', ' ', txt).strip()
        
        if txt and len(txt) > 50:  # Good extraction
            return txt
            
    except Exception as e:
        logger.warning(f"Mammoth extraction failed: {e}")
    
    # Fallback to python-docx
    try:
        doc = Document(path)
        parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        
        return "\n".join(parts)
        
    except Exception as e:
        logger.error(f"python-docx extraction failed: {e}")
        return ""

def extract_text_safely(path: str, expect_arabic=True) -> str:
    """Extract text with tiered approach and OCR fallback."""
    ext = os.path.splitext(path.lower())[1]
    text = ""
    
    logger.info(f"Extracting text from {path} (expect_arabic={expect_arabic})")
    
    if ext == ".docx":
        # Extract from DOCX
        text = _read_docx(path)
        logger.info(f"DOCX extraction: {len(text)} chars, Arabic ratio: {len(ARABIC_RE.findall(text))/max(1,len(text)):.2f}")
        
        # If expecting Arabic but not found, try OCR via PDF conversion
        if expect_arabic and text and not contains_arabic(text):
            logger.info("Arabic expected but not found, trying OCR...")
            pdf_path = _convert_to_pdf(path)
            if pdf_path:
                ocr_text = _ocr_pdf(pdf_path)
                if ocr_text and contains_arabic(ocr_text):
                    logger.info(f"OCR success: {len(ocr_text)} chars")
                    text = ocr_text
                
    elif ext == ".doc":
        # Convert old DOC to PDF first
        pdf_path = _convert_to_pdf(path)
        if pdf_path:
            text = _pdf_text(pdf_path)
            
            # OCR fallback if needed
            if expect_arabic and text and not contains_arabic(text):
                ocr_text = _ocr_pdf(pdf_path)
                if ocr_text and contains_arabic(ocr_text):
                    text = ocr_text
                    
    elif ext == ".pdf":
        # Extract from PDF
        text = _pdf_text(path)
        logger.info(f"PDF extraction: {len(text)} chars")
        
        # OCR fallback if Arabic expected but not found
        if expect_arabic and text and not contains_arabic(text):
            logger.info("Arabic expected but not found in PDF, trying OCR...")
            ocr_text = _ocr_pdf(path)
            if ocr_text and contains_arabic(ocr_text):
                logger.info(f"OCR success: {len(ocr_text)} chars")
                text = ocr_text
    else:
        # Handle other text files
        try:
            with open(path, "rb") as f:
                raw = f.read()
            try:
                text = raw.decode("utf-8")
            except UnicodeDecodeError:
                # Fallback to Windows-1256 for Arabic
                text = raw.decode("windows-1256", errors="ignore")
        except Exception as e:
            logger.error(f"Text file reading failed: {e}")
            return ""
    
    # Normalize Arabic text
    if text:
        text = normalize_arabic(text)
        logger.info(f"Final text: {len(text)} chars after normalization")
    
    return text

def extraction_health(s: str) -> dict:
    """Calculate health metrics for extracted text."""
    if not s: 
        return {"length": 0, "arabic_ratio": 0, "health": "empty"}
    
    arab_count = len(ARABIC_RE.findall(s))
    arabic_ratio = arab_count / len(s)
    
    # Determine health status
    if arabic_ratio >= 0.15:
        health = "good"
    elif arabic_ratio >= 0.05:
        health = "suspect"
    else:
        health = "poor"
    
    return {
        "length": len(s), 
        "arabic_ratio": arabic_ratio,
        "arabic_chars": arab_count,
        "health": health
    }