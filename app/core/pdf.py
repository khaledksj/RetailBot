"""
Document processing utilities for the Shop Manual Chatbot RAG system.
Supports PDF and Word (.docx) documents.
"""

import io
from typing import List, Optional
import pypdf
from pypdf import PdfReader
from pdfminer.high_level import extract_text
from docx import Document

from app.core.logging import get_logger

logger = get_logger(__name__)

class DocumentProcessor:
    """Document processor supporting PDF and Word (.docx) files with multiple extraction backends."""
    
    async def extract_text(self, document_content: bytes, file_type: str = "pdf") -> List[str]:
        """
        Extract text from document content (PDF or Word).
        
        Args:
            document_content: Document file content as bytes
            file_type: Type of document ("pdf" or "docx")
            
        Returns:
            List of text strings, one per page (PDF) or one per paragraph (Word)
        """
        if file_type.lower() == "docx":
            return await self._extract_from_word(document_content)
        else:
            return await self._extract_from_pdf(document_content)
    
    async def _extract_from_pdf(self, pdf_content: bytes) -> List[str]:
        """Extract text from PDF content."""
        logger.info(f"Extracting text from PDF", extra={
            "pdf_size_bytes": len(pdf_content)
        })
        
        # Try pypdf first
        try:
            pages_text = await self._extract_with_pypdf(pdf_content)
            if pages_text and any(page.strip() for page in pages_text):
                logger.info(f"Successfully extracted text with pypdf", extra={
                    "pages_count": len(pages_text),
                    "total_chars": sum(len(page) for page in pages_text)
                })
                return pages_text
        except Exception as e:
            logger.warning(f"pypdf extraction failed: {str(e)}")
        
        # Fallback to pdfminer
        try:
            pages_text = await self._extract_with_pdfminer(pdf_content)
            if pages_text and any(page.strip() for page in pages_text):
                logger.info(f"Successfully extracted text with pdfminer", extra={
                    "pages_count": len(pages_text),
                    "total_chars": sum(len(page) for page in pages_text)
                })
                return pages_text
        except Exception as e:
            logger.warning(f"pdfminer extraction failed: {str(e)}")
        
        logger.error("All PDF extraction methods failed")
        return []
    
    async def _extract_from_word(self, docx_content: bytes) -> List[str]:
        """Extract text from Word (.docx) document."""
        try:
            logger.info(f"Extracting text from Word document", extra={
                "docx_size_bytes": len(docx_content)
            })
            
            # Load the Word document from bytes
            doc_file = io.BytesIO(docx_content)
            doc = Document(doc_file)
            
            # Extract text paragraph by paragraph
            paragraphs = []
            current_page_text = []
            chars_per_page = 3000  # Approximate characters per "page" for chunking
            current_chars = 0
            
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:  # Skip empty paragraphs
                    current_page_text.append(para_text)
                    current_chars += len(para_text)
                    
                    # If we've accumulated enough text, create a new "page"
                    if current_chars >= chars_per_page:
                        paragraphs.append('\n\n'.join(current_page_text))
                        current_page_text = []
                        current_chars = 0
            
            # Add any remaining text as the final "page"
            if current_page_text:
                paragraphs.append('\n\n'.join(current_page_text))
            
            # If no paragraphs were found, try extracting from tables
            if not paragraphs:
                table_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            table_text.append(' | '.join(row_text))
                
                if table_text:
                    paragraphs = ['\n'.join(table_text)]
            
            logger.info(f"Successfully extracted text from Word document", extra={
                "pages_count": len(paragraphs),
                "total_chars": sum(len(page) for page in paragraphs)
            })
            
            return paragraphs if paragraphs else [""]
            
        except Exception as e:
            logger.error(f"Failed to extract text from Word document: {str(e)}")
            return []
    
    async def _extract_with_pypdf(self, pdf_content: bytes) -> List[str]:
        """Extract text using pypdf library with Arabic support."""
        pdf_file = io.BytesIO(pdf_content)
        reader = PdfReader(pdf_file)
        
        pages_text = []
        for page_num, page in enumerate(reader.pages):
            try:
                # Extract text with different strategies for better Arabic support
                text = page.extract_text()
                
                # If text extraction results in mostly dots or symbols, try alternative method
                if text and self._is_corrupted_text(text):
                    logger.warning(f"Detected corrupted text on page {page_num + 1}, trying alternative extraction")
                    # Skip pypdf alternative for now - it's not working well
                    # The pdfminer fallback should handle this better
                    pass
                
                pages_text.append(text or "")
            except Exception as e:
                logger.warning(f"Error extracting page {page_num + 1}: {str(e)}")
                pages_text.append("")
        
        return pages_text
    
    def _is_corrupted_text(self, text: str) -> bool:
        """Check if extracted text appears to be corrupted (mostly dots, symbols)."""
        if not text:
            return True
        
        # Count meaningful characters vs dots/symbols
        meaningful_chars = sum(1 for c in text if c.isalnum())
        dots_and_symbols = sum(1 for c in text if c in '.,;:!?-_()[]{}')
        total_chars = len(text.replace(' ', '').replace('\n', ''))
        
        if total_chars == 0:
            return True
        
        # If more than 60% are dots/symbols, consider it corrupted
        corruption_ratio = dots_and_symbols / total_chars
        return corruption_ratio > 0.6
    
    async def _extract_with_pdfminer(self, pdf_content: bytes) -> List[str]:
        """Extract text using pdfminer library with better Arabic support."""
        pdf_file = io.BytesIO(pdf_content)
        
        # Extract all text with better encoding handling
        try:
            # Use LAParams for better Arabic text extraction
            from pdfminer.layout import LAParams
            from pdfminer.converter import TextConverter
            from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
            from pdfminer.pdfpage import PDFPage
            
            output_string = io.StringIO()
            laparams = LAParams(
                char_margin=0.3,
                line_margin=0.5,
                word_margin=0.1,
                boxes_flow=0.5,
                detect_vertical=True  # Better for Arabic text
            )
            
            with TextConverter(PDFResourceManager(), output_string, laparams=laparams) as device:
                interpreter = PDFPageInterpreter(PDFResourceManager(), device)
                for page in PDFPage.get_pages(pdf_file):
                    interpreter.process_page(page)
            
            full_text = output_string.getvalue()
            output_string.close()
            
        except Exception as e:
            logger.warning(f"Advanced pdfminer extraction failed: {str(e)}")
            # Fallback to simple extraction
            try:
                pdf_file.seek(0)  # Reset file pointer
                full_text = extract_text(pdf_file)
            except Exception as e2:
                logger.error(f"pdfminer full text extraction failed: {str(e2)}")
                return []
        
        # Simple heuristic to split by pages
        pages = self._split_text_by_pages(full_text)
        
        return pages
    
    def _split_text_by_pages(self, text: str) -> List[str]:
        """
        Split extracted text into pages using heuristics.
        This is a simple approach and may not work perfectly for all PDFs.
        """
        # Common page break indicators
        page_break_patterns = [
            '\f',  # Form feed character
            '\n\n\n',  # Multiple newlines
        ]
        
        # Try form feed first (most reliable)
        if '\f' in text:
            pages = text.split('\f')
            return [page.strip() for page in pages if page.strip()]
        
        # Fallback: split by multiple newlines and try to identify page boundaries
        paragraphs = text.split('\n\n')
        
        # If we have a reasonable number of paragraphs, group them
        if len(paragraphs) > 5:
            # Simple heuristic: assume roughly equal distribution
            pages_count = max(1, len(paragraphs) // 10)  # Assume ~10 paragraphs per page
            page_size = len(paragraphs) // pages_count
            
            pages = []
            for i in range(0, len(paragraphs), page_size):
                page_paragraphs = paragraphs[i:i + page_size]
                page_text = '\n\n'.join(page_paragraphs)
                if page_text.strip():
                    pages.append(page_text.strip())
            
            return pages
        
        # Last resort: return as single page
        return [text.strip()] if text.strip() else []
    
    def validate_pdf(self, pdf_content: bytes) -> bool:
        """Validate that the content is a valid PDF."""
        try:
            pdf_file = io.BytesIO(pdf_content)
            reader = PdfReader(pdf_file)
            
            # Check if we can read at least the number of pages
            page_count = len(reader.pages)
            return page_count > 0
            
        except Exception as e:
            logger.warning(f"PDF validation failed: {str(e)}")
            return False
